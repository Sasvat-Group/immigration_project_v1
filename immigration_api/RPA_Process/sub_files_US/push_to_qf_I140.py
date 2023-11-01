import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx2pdf import convert
import os
import os.path as op
import json
import shutil
import glob
import requests
import base64
from datetime import date
from pathlib import Path


def generate_140(ProjectID, formtype, other_objects):
    mailing_address = other_objects['mailing_address']
    current_res_address = other_objects['current_res_address']
    last_perm_address_abroad = other_objects['last_perm_address_abroad']
    perm_foreign_address = other_objects['perm_foreign_address']
    ben_cur_emp = other_objects['ben_cur_emp']
    bit_ans = other_objects['bit_ans']
    iff = other_objects['iff']
    result = other_objects['result']

    most_recent_entry_rec = other_objects['most_recent_entry_rec']
    most_recent_notice = other_objects['most_recent_notice']

    conn = other_objects['conn']
    cursor = conn.cursor()
    headersAPI = other_objects['headersAPI']
    qpdf_merger = other_objects['qpdf_merger']
    parent_dir = other_objects['parent_dir']
    change_date = other_objects['change_date']
    AptSteFlr = other_objects['AptSteFlr']
    iff = other_objects['iff']
    return_high_salary = other_objects['return_high_salary']
    results2 = other_objects['results2']

    if result:
        if True:
            dependents = cursor.execute(
                '''select *,DATEDIFF(YY,BirthDate,GETDATE()) as Age from [dbo].[Beneficiary] where PrimaryBeneficiaryID = '{}' and not (LOWER(BirthCountry) Like '%america%' or LOWER(BirthCountry) Like '%u.s.a%' or LOWER(BirthCountry) Like '%usa%' ) and not (LOWER(CitizenshipCountry) Like '%america%' or LOWER(CitizenshipCountry) Like '%u.s.a%' or LOWER(CitizenshipCountry) Like '%usa%' ) and ((LOWER(RelationType) in('husband','spouse','wife'))  or (lower(RelationType) in('child','son','daughter','biological son','biological daughter','biological child') and ((DATEDIFF(YY,BirthDate,GETDATE()))<21))) '''.format(
                    result.BeneficiaryID)).fetchall()
            cursor.close()
            p1 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}
            p2 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}
            p3 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}
            p4 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}
            p5 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}
            p6 = {'FirstName': '', 'LastName': '', 'MiddleName': '',
                  'DateOfBirth': '', 'BirthCountry': '', 'Relation': ''}

            persons = [p1, p2, p3, p4, p5, p6]

            for dependent, person in zip(dependents, persons):
                person['FirstName'] = dependent.FirstName
                person['LastName'] = dependent.LastName
                person['MiddleName'] = dependent.MiddleName
                person['DateOfBirth'] = dependent.BirthDate
                person['BirthCountry'] = dependent.BirthCountry
                person['Relation'] = dependent.RelationType

            result_priority_category_chkbox_val = ''
            if str(result.ProjectPriorityCategory) in ["EB2", "EB-2", "Employment-Based 2nd", "Employment Based 2nd",
                                                       "203(b)(2)"]:
                result_priority_category_chkbox_val = 4
            if str(result.ProjectPriorityCategory) in ["EB3", "EB-3", "Employment-Based 3rd", "Employment Based 3rd",
                                                       "203(b)(3)(A)(ii)"]:
                result_priority_category_chkbox_val = 5

            part_4_checkbox_text = perm_foreign_address.AddressUnitType if str(
                mailing_address.Country).lower() in ('usa', 'u.s.a', 'america', 'united states of america') else ''

            data_dict = [
                {
                    "FieldName": "USIMM82474.txt3.178",  # p42b
                    "FieldValue": current_res_address.Country if str(current_res_address.Country).lower().strip() not in ('usa', 'u.s.a', 'america', 'united states of america') else last_perm_address_abroad.Country
                },
                {
                    "FieldName": "1ben.F.Addr123",  # p43a
                    "FieldCalcOverride": True,
                    "FieldValue": perm_foreign_address.Address1 if str(mailing_address.Country).lower().strip() in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.F.AddrUnitType",  # p43b checkbox
                    "FieldValue": AptSteFlr(part_4_checkbox_text)
                },
                {
                    "FieldName": "1ben.F.Addr4",  # p43b blankbx
                    "FieldValue": perm_foreign_address.AddressUnitNumber if str(mailing_address.Country).lower().strip() in ('usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.F.City",  # p43c blankbx
                    "FieldValue": perm_foreign_address.City if str(mailing_address.Country).lower().strip() in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.F.County",  # p43d blankbx
                    "FieldValue": perm_foreign_address.StateProvince if str(
                        mailing_address.Country).lower().strip() in ('usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.F.ForeignRoute",  # p43e blankbx
                    "FieldValue": perm_foreign_address.PostalZipCode if str(
                        mailing_address.Country).lower().strip() in ('usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.F.Country",  # p43f blankbx
                    "FieldValue": perm_foreign_address.Country if str(mailing_address.Country).lower().strip() in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "3altclient.M.Addr123",
                    "FieldCalcOverride": True,
                    "FieldValue": iff(results2.WorkLocationAddress1)
                },
                {
                    "FieldName": "3altclient.M.Addr4",
                    "FieldValue": iff(results2.WorkLocationAddress2)
                },
                {
                    "FieldName": "3altclient.M.City",
                    "FieldValue": iff(results2.WorkLocationCity)
                },
                {
                    "FieldName": "3altclient.M.State",
                    "FieldValue": iff(results2.WorkLocationState)
                },
                {
                    "FieldName": "3altclient.M.Zip",
                    "FieldValue": iff(results2.WorkLocationZipCode)
                },
                {
                    "FieldName": "1spou.FName",
                    "FieldValue": iff(p1['FirstName'])
                },
                {
                    "FieldName": "1spou.LName",
                    "FieldValue": iff(p1['LastName'])
                },
                {
                    "FieldName": "1spou.MName",
                    "FieldValue": iff(p1['MiddleName'])
                },
                {
                    "FieldName": "1spou.DOB",
                    "FieldValue": change_date(p1['DateOfBirth'])
                },
                {
                    "FieldName": "1spou.BP.Country",
                    "FieldValue": iff(p1['BirthCountry'])
                },
                {
                    "FieldName": "1spou.RelToClient",
                    "FieldValue": iff(p1['Relation'])
                },
                {
                    "FieldName": "1child.FName",
                    "FieldValue": iff(p2['FirstName'])
                },
                {
                    "FieldName": "1child.LName",
                    "FieldValue": iff(p2['LastName'])
                },
                {
                    "FieldName": "1child.MName",
                    "FieldValue": iff(p2['MiddleName'])
                },
                {
                    "FieldName": "1child.DOB",
                    "FieldValue": change_date(p2['DateOfBirth'])
                },
                {
                    "FieldName": "1child.BP.Country",
                    "FieldValue": iff(p2['BirthCountry'])
                },
                {
                    "FieldName": "1child.RelToClient",
                    "FieldValue": iff(p2['Relation'])
                },
                {
                    "FieldName": "2child.FName",
                    "FieldValue": iff(p3['FirstName'])
                },
                {
                    "FieldName": "2child.LName",
                    "FieldValue": iff(p3['LastName'])
                },
                {
                    "FieldName": "2child.MName",
                    "FieldValue": iff(p3['MiddleName'])
                },
                {
                    "FieldName": "2child.DOB",
                    "FieldValue": change_date(p3['DateOfBirth'])
                },
                {
                    "FieldName": "2child.BP.Country",
                    "FieldValue": iff(p3['BirthCountry'])
                },
                {
                    "FieldName": "2child.RelToClient",
                    "FieldValue": iff(p3['Relation'])
                },
                {
                    "FieldName": "3child.FName",
                    "FieldValue": iff(p4['FirstName'])
                },
                {
                    "FieldName": "3child.LName",
                    "FieldValue": iff(p4['LastName'])
                },
                {
                    "FieldName": "3child.MName",
                    "FieldValue": iff(p4['MiddleName'])
                },
                {
                    "FieldName": "3child.DOB",
                    "FieldValue": change_date(p4['DateOfBirth'])
                },
                {
                    "FieldName": "3child.BP.Country",
                    "FieldValue": iff(p4['BirthCountry'])
                },
                {
                    "FieldName": "3child.RelToClient",
                    "FieldValue": iff(p4['Relation'])
                },
                {
                    "FieldName": "4child.FName",
                    "FieldValue": iff(p5['FirstName'])
                },
                {
                    "FieldName": "4child.LName",
                    "FieldValue": iff(p5['LastName'])
                },
                {
                    "FieldName": "4child.MName",
                    "FieldValue": iff(p5['MiddleName'])
                },
                {
                    "FieldName": "4child.DOB",
                    "FieldValue": change_date(p5['DateOfBirth'])
                },
                {
                    "FieldName": "4child.BP.Country",
                    "FieldValue": iff(p5['BirthCountry'])
                },
                {
                    "FieldName": "4child.RelToClient",
                    "FieldValue": iff(p5['Relation'])
                },
                {
                    "FieldName": "5child.FName",
                    "FieldValue": iff(p6['FirstName'])
                },
                {
                    "FieldName": "5child.LName",
                    "FieldValue": iff(p6['LastName'])
                },
                {
                    "FieldName": "5child.MName",
                    "FieldValue": iff(p6['MiddleName'])
                },
                {
                    "FieldName": "5child.DOB",
                    "FieldValue": change_date(p6['DateOfBirth'])
                },
                {
                    "FieldName": "5child.BP.Country",
                    "FieldValue": iff(p6['BirthCountry'])
                },
                {
                    "FieldName": "5child.RelToClient",
                    "FieldValue": iff(p6['Relation'])
                },
                {
                    "FieldName": "1lawyer.G28Attached",
                    "FieldValue": 1
                },
                {
                    "FieldName": "USIMM82474.txt3.83",
                    "FieldValue": iff(results2.OfferedWageType)
                },
                {
                    "FieldName": "USIMM82474.txt3.79",
                    "FieldValue": return_high_salary(results2.OfferedWageFrom, ben_cur_emp.CurrentAnnualSalary)},
                {
                    "FieldName": "USIMM82474.txt3.74",
                    "FieldValue": iff(results2.JobTitle)
                },
                {
                    "FieldName": "USIMM82474.txt3.78.0",
                    "FieldValue": iff(results2.SOCCODE.split('-')[0].strip())
                },
                {
                    "FieldName": "USIMM82474.txt3.78.1",
                    "FieldValue": iff(results2.SOCCODE.split('-')[1].strip())
                },
                {
                    "FieldName": "USIMM82474.txt3.68.0",
                    "FieldValue": change_date(results2.PERMFilingDate)
                },
                {
                    "FieldName": "USIMM82474.txt3.68.1",
                    "FieldValue": change_date(results2.PERMValidTo)
                },
                {
                    "FieldName": "USIMM82474.txt3.38",
                    # if results.ReceiptNumber #is not None else results2.PERMDOLProjectNumber
                    "FieldValue": iff(result.ReceiptNumber)
                },
                {
                    "FieldName": "1lawyer.LicNum",
                    "FieldValue": iff(result.l_BarNumber)
                },
                {
                    "FieldName": "1ben.LName",
                    "FieldValue": result.LastName if result.LastName is not None else results2.PetitionerContactLastName
                },
                {
                    "FieldName": "1ben.FName",
                    "FieldValue": result.FirstName if result.FirstName is not None else results2.PetitionerContactFirstName
                },
                {
                    "FieldName": "1ben.MName",
                    "FieldValue": result.MiddleName if result.MiddleName is not None else results2.PetitionerContactMiddleName
                },
                {
                    "FieldName": "1contact.SEmp.Company",
                    "FieldValue": result.PetitionerName if result.PetitionerName is not None else (str(results2.PetitionerContactLastName) + ', ' + str(results2.PetitionerContactFirstName) + ' ' + str(results2.PetitionerContactMiddleName)).strip()
                },
                {
                    "FieldName": "1contact.EO.Attn",
                    "FieldValue": (str(result.FirstName if result.FirstName is not None else results2.PetitionerContactFirstName) + ' ' + str(result.LastName if result.LastName is not None else results2.PetitionerContactLastName)).strip()
                },
                {
                    "FieldName": "txtCalcEO.Addr1",
                    "FieldValue": result.Address1 if result.Address1 is not None else results2.PetitionerContactAddress1
                },
                {
                    "FieldName": "USIMM82474.chk3.11.6",
                    "FieldValue": 'No'
                },
                {
                    "FieldName": "1ben.PendingProceedings",
                    "FieldValue": bit_ans(result.IsInRemovalProceeding)
                },
                {  # part4 2a
                    "FieldName": "USIMM82474.chk3.07",
                    "FieldValue": '2'
                },
                {  # part4 6a
                    "FieldName": "1ben.FilingOtherPetitions",
                    "FieldValue": 'No'
                },
                {
                    "FieldName": "1contact.EO.Addr123",  # overwri 1contact.EO.Addr123
                    "FieldCalcOverride": True,
                    "FieldValue": iff(result.Address1)
                },
                {  # peculiar
                    "FieldName": "1contact.EO.Addr4",
                    "FieldValue": result.AddressUnitNumber if result.AddressUnitNumber is not None else results2.PetitonerContactAddress2
                },
                {
                    "FieldName": "1contact.EO.AddrUnitType",
                    "FieldValue": AptSteFlr(result.AddressUnitType)
                },
                {
                    "FieldName": "User.D937.AttorneyOrAccreditedRep",
                    "FieldValue": '1'
                },
                {
                    "FieldName": "USIMM82474.chk3.27.0",
                    "FieldValue": 'No' if p1['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk3.26.0",
                    "FieldValue": 'Yes' if p1['FirstName'] else ''
                },
                {
                    "FieldName": "1ben.PreparerPreparedApp",
                    "FieldValue": '1'
                },
                {
                    "FieldName": "USIMM82474.chk3.26.1",
                    "FieldValue": 'Yes' if p2['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk3.27.1",
                    "FieldValue": 'No' if p2['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.26.0.0",
                    "FieldValue": 'Yes' if p3['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.27.0.0",
                    "FieldValue": 'No' if p3['FirstName'] else ''
                },
                {
                    "FieldName": "1lawyer.ExtendsPrepApp",
                    "FieldValue": '1'
                },
                {
                    "FieldName": "USIMM82474.chk2.26.0.1",
                    "FieldValue": 'Yes' if p4['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.27.0.1",
                    "FieldValue": 'No' if p4['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.26.1.0",
                    "FieldValue": 'Yes' if p5['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.27.1.0",
                    "FieldValue": 'No' if p5['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.26.1.1",
                    "FieldValue": 'Yes' if p6['FirstName'] else ''
                },
                {
                    "FieldName": "USIMM82474.chk2.27.1.1",
                    "FieldValue": 'No' if p6['FirstName'] else ''
                },
                {
                    "FieldName": "1ben.UnderstandLanguage",
                    "FieldValue": '1'
                },
                {
                    "FieldName": "User.D937.NontechnicalDescriptionOfJob.multiline.1",
                    "FieldValue": "See attached Certified ETA Form 9089"
                },
                {
                    "FieldName": "1contact.EO.City",
                    "FieldValue": result.City if result.City is not None else results2.PetitionerContactCity
                },
                {
                    "FieldName": "USIMM82474.chk3.22",
                    "FieldValue": '1'
                },
                {
                    "FieldName": "USIMM82474.chk3.24",
                    "FieldValue": 'Yes'
                },
                {
                    "FieldName": "USIMM82474.chk3.23.0",
                    "FieldValue": 'Yes'
                },
                {
                    "FieldName": "1contact.EO.State",
                    "FieldValue": result.StateProvince if result.StateProvince is not None else results2.PetitionerContactState
                },
                {
                    "FieldName": "1contact.EO.Zip",
                    "FieldValue": result.PostalZipCode if result.PostalZipCode is not None else results2.PetitionerContactZipCode
                },
                {
                    "FieldName": "1contact.EO.Country",
                    "FieldValue": result.Country if result.Country is not None else results2.PetitionerContactCountry
                },
                {
                    "FieldName": "1contact.SEmp.TaxIDComb",
                    "FieldValue": result.FederalEmployerId.replace('-', '') if result.FederalEmployerId is not None else str(results2.FEIN).replace('-', '')
                },
                {
                    "FieldName": "USIMM82474.chk1.04",
                    "FieldValue": result_priority_category_chkbox_val
                },
                {
                    "FieldName": "1ben.LName",
                    "FieldValue": result.b_lastname if result.b_lastname is not None else results2.BeneficiaryLastName
                },
                {
                    "FieldName": "1ben.FName",
                    "FieldValue": result.b_firstname if result.b_firstname is not None else results2.BeneficiaryFirstName
                },
                {
                    "FieldName": "1ben.MName",
                    "FieldValue": result.b_middlename if result.b_middlename is not None else results2.BeneficiaryMiddleName
                },
                {
                    "FieldName": "1ben.M.Attn",
                    "FieldValue": (str(result.b_firstname if result.b_firstname is not None else results2.BeneficiaryFirstName) + ' ' + str(result.b_lastname if result.b_lastname is not None else results2.BeneficiaryLastName)).strip()
                },
                {
                    "FieldName": "1ben.M.Addr123",
                    "FieldCalcOverride": True,
                    "FieldValue": mailing_address.Address1 if mailing_address.Address1 is not None else results2.BeneficiaryAddress1
                },
                {
                    "FieldName": "1ben.M.AddrUnitType",
                    "FieldValue": AptSteFlr(mailing_address.AddressUnitType)
                },
                {
                    "FieldName": "1ben.M.Addr4",
                    "FieldValue": mailing_address.AddressUnitNumber if mailing_address.AddressUnitNumber is not None else results2.BeneficiaryAddress2
                },
                {
                    "FieldName": "1ben.M.City",
                    "FieldValue": mailing_address.City if mailing_address.City is not None else results2.BeneficiaryAddressCity
                },
                {
                    "FieldName": "1ben.M.State",
                    "FieldValue": mailing_address.StateProvince if str(mailing_address.Country).lower().strip() in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.M.County",
                    "FieldValue": mailing_address.StateProvince if str(mailing_address.Country).lower().strip() not in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.M.Zip",
                    "FieldValue": mailing_address.PostalZipCode if str(mailing_address.Country).lower().strip() in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.M.ForeignRoute",
                    "FieldValue": mailing_address.PostalZipCode if str(mailing_address.Country).lower().strip() not in (
                    'usa', 'u.s.a', 'america', 'united states of america') else ''
                },
                {
                    "FieldName": "1ben.M.Country",
                    "FieldValue": mailing_address.Country if mailing_address.Country is not None else results2.BeneficiaryAddressCountry
                },
                {
                    "FieldName": "1ben.DOB",
                    "FieldValue": result.BirthDate.strftime('%m/%d/%Y') if result.BirthDate is not None else results2.BeneficiaryDateofBirth.strftime('%m/%d/%Y')
                },
                {
                    "FieldName": "1ben.BP.City",
                    "FieldValue": iff(result.BirthCity)
                },
                {
                    "FieldName": "1ben.BP.State",
                    "FieldValue": iff(result.BirthStateProvince)
                },
                {
                    "FieldName": "1ben.BP.Country",
                    "FieldValue": result.BirthCountry if result.BirthCountry is not None else results2.BeneficiaryCountryofBirth
                },
                {
                    "FieldName": "1ben.Citizenship",
                    "FieldValue": result.CitizenshipCountry if result.CitizenshipCountry is not None else results2.BeneficiaryCountryofCitizenship
                },
                {
                    "FieldName": "1ben.AlienRegNumComb",
                    "FieldValue": result.AlienNumber1 if result.AlienNumber1 is not None else results2.BeneficiaryAlienNumber
                },
                {
                    "FieldName": "1ben.SSNComb",
                    "FieldValue": iff(result.SSNNumber)
                },
                {
                    "FieldName": "1ben.I94.ArrivalDate",
                    "FieldValue": change_date(most_recent_entry_rec.EntryDate)
                },
                {
                    "FieldName": "1ben.I94.IDComb",
                    "FieldValue": iff(most_recent_entry_rec.I94Number)
                },
                {
                    "FieldName": "1ben.I94.IDExpire",
                    "FieldValue": change_date(most_recent_entry_rec.I94ExpirationDate)
                },
                {
                    "FieldName": "1ben.I94.IDStatus",
                    "FieldValue": iff(most_recent_entry_rec.I94StatusType)
                },
                {
                    "FieldName": "1ben.GOV.ID",
                    "FieldValue": iff(most_recent_entry_rec.TravelDocumentNumber)
                },
                {
                    "FieldName": "1ben.GOV.IDCountry",
                    "FieldValue": iff(most_recent_entry_rec.IssuingCountry)
                },
                {
                    "FieldName": "1ben.GOV.IDExpire",
                    "FieldValue": change_date(most_recent_entry_rec.ValidTo)
                },
                {
                    "FieldName": "USIMM82474.txt3.19.1",
                    "FieldValue": (str(result.b_firstname if result.b_firstname is not None else results2.BeneficiaryFirstName) + ' ' + str(result.b_lastname if result.b_lastname is not None else results2.BeneficiaryLastName)).strip()
                },
                {
                    "FieldName": "USIMM82474.txt3.19.0",
                    "FieldValue": mailing_address.Address1 if mailing_address.Address1 is not None else results2.BeneficiaryAddress1
                },
                {
                    "FieldName": "USIMM82474.txt3.24.0",
                    "FieldValue": mailing_address.AddressUnitNumber if mailing_address.AddressUnitNumber is not None else results2.BeneficiaryAddress2
                },
                {
                    "FieldName": "USIMM82474.chk3.09",
                    "FieldValue": AptSteFlr(mailing_address.AddressUnitType)
                },
                {
                    "FieldName": "USIMM82474.txt3.29.0",
                    "FieldValue": mailing_address.City if mailing_address.City is not None else results2.BeneficiaryAddressCity
                },
                {
                    "FieldName": "USIMM82474.txt3.29.2",
                    "FieldValue": mailing_address.StateProvince if mailing_address.StateProvince is not None else results2.BeneficiaryAddressState
                },
                {
                    "FieldName": "USIMM82474.txt3.29.3",
                    "FieldValue": mailing_address.PostalZipCode if mailing_address.PostalZipCode is not None else results2.BeneficiaryAddressZipCode
                },
                {
                    "FieldName": "USIMM82474.txt3.34",
                    "FieldValue": mailing_address.Country if mailing_address.Country is not None else results2.BeneficiaryAddressCountry
                },
                {
                    "FieldName": "1contact.SEmp.BusType",
                    "FieldValue": iff(result.p_BusinessType)
                },
                {
                    "FieldName": "1contact.SEmp.YearEst",
                    "FieldValue": result.p_YearEstablished if result.p_YearEstablished is not None else results2.YearCommencedBusiness
                },
                {
                    "FieldName": "User.D937.1contact.SEmp.NumEmployeesUS",
                    "FieldValue": result.p_EmployeeCountUS if result.p_EmployeeCountUS is not None else results2.NumberOfEmployees
                },
                {
                    "FieldName": "1contact.SEmp.GrossIncome",
                    "FieldValue": iff(result.p_AnnualIncomeGrossUS)
                },
                {
                    "FieldName": "1contact.SEmp.NetIncome",
                    "FieldValue": iff(result.p_AnnualIncomeNetUS)
                },
                {
                    "FieldName": "1contact.SEmp.NAICSComb",
                    "FieldValue": result.p_NaicsCode if result.p_NaicsCode is not None else results2.NAICSCode
                },
                {
                    "FieldName": "1laywer.FullName",
                    "FieldValue": iff((str(result.l_firstname) + ' ' + str(result.l_lastname)).strip())
                },
                {
                    "FieldName": "1authind.LName",
                    "FieldValue": result.LastName if result.LastName is not None else results2.PetitionerContactLastName
                },
                {
                    "FieldName": "1authind.FName",
                    "FieldValue": result.FirstName if result.FirstName is not None else results2.PetitionerContactFirstName
                },
                {
                    "FieldName": "1authind.SignTitle",
                    "FieldValue": result.JobTitle if result.JobTitle is not None else results2.JobTitle
                },
                {
                    "FieldName": "1authind.O.Phone",
                    "FieldValue": result.PhoneNumber if result.PhoneNumber is not None else results2.PetitionerContactPhoneNumber
                },
                {
                    "FieldName": "USIMM82474.txt3.525",
                    "FieldValue": iff(result.MobilePhone)
                },
                {
                    "FieldName": "1authind.O.Email",
                    "FieldValue": result.Email if result.Email is not None else results2.PetitionerContactEmail
                },
                {
                    "FieldName": "1lawyer.LName",
                    "FieldValue": result.l_lastname if result.l_lastname is not None else results2.AttorneyAgentLastName
                },
                {
                    "FieldName": "1lawyer.FName",
                    "FieldValue": result.l_firstname if result.l_firstname is not None else results2.AttorneyAgentFirstName
                },
                {
                    "FieldName": "1lawfirm.Company",
                    "FieldValue": result.l_firmname if result.l_firmname is not None else results2.AttorneyAgentFirmName
                },
                {
                    "FieldName": "1lawfirm.O.Addr123",
                    "FieldCalcOverride": True,
                    "FieldValue": result.l_address1 if result.l_address1 is not None else results2.AttorneyAgentAddress1
                },
                {
                    "FieldName": "1lawfirm.O.AddrUnitType",
                    "FieldValue": AptSteFlr(result.l_addressunittype)
                },
                {
                    "FieldName": "1lawfirm.O.Addr4",
                    "FieldValue": result.l_addressunitnumber if result.l_addressunitnumber is not None else results2.AttorneyAgentAddress2
                },
                {
                    "FieldName": "1lawfirm.O.City",
                    "FieldValue": result.l_city if result.l_city is not None else results2.AttorneyAgentCity
                },
                {
                    "FieldName": "1lawfirm.O.State",
                    "FieldValue": result.l_state if result.l_state is not None else results2.AttorneyAgentState
                },
                {
                    "FieldName": "1lawfirm.O.Zip",
                    "FieldValue": result.l_zipcode if result.l_zipcode is not None else results2.AttorneyAgentZipCode
                },
                {
                    "FieldName": "1lawfirm.O.Country",
                    "FieldValue": result.l_country if result.l_country is not None else results2.AttorneyAgentCountry
                },
                {
                    "FieldName": "1lawyer.O.Phone",
                    "FieldValue": result.l_phonenumber if result.l_phonenumber is not None else results2.AttorneyAgentPhoneNumber
                },
                {
                    "FieldName": "1lawyer.O.Mobile",
                    "FieldValue": iff(result.l_mobilenumber)
                },
                {
                    "FieldName": "1lawyer.O.Email",
                    "FieldValue": result.l_email if result.l_email is not None else results2.AttorneyAgentEmail
                },
                {
                    "FieldName": "1lawyer.ELISNumComb",
                    "FieldValue": iff(result.l_USCISNumber)
                },
                {
                    "FieldName": "1contact.SEmp.SSNComb",
                    "FieldValue": iff(result.SSNNumber)
                },
                {
                    "FieldName": "USIMM82474.txt1.08",
                    "FieldValue": iff(result.b_USCISNumber)
                }
            ]

            params = {
                "HostFormOnQuik": True,
                "FormFields": data_dict,
                "QuikFormID": "82474",
                "PrintEditablePDF": True
            }

            data_json = json.dumps(params)
            response = requests.post('https://websvcs.quikforms.com/rest/quikformsengine/qfe/execute/pdf',
                                     headers=headersAPI, data=data_json)
            api_response = response.json()
            pdf_base64 = api_response['ResultData']['PDF']
            with open('Form I-140.pdf', 'wb') as pdf:
                pdf.write(base64.b64decode(pdf_base64))

            data_dict2 = [
                {
                    "FieldName": "1ben.AlienRegNumComb",
                    "FieldValue": result.AlienNumber1 if result.AlienNumber1 is not None else results2.BeneficiaryAlienNumber
                },
                {
                    "FieldName": "1lawyer.ELISNumComb",
                    "FieldValue": iff(result.l_USCISNumber)
                },
                {
                    "FieldName": "1lawyer.LName",
                    "FieldValue": result.l_lastname if result.l_lastname is not None else results2.AttorneyAgentLastName
                },
                {
                    "FieldName": "1lawyer.FName",
                    "FieldValue": result.l_firstname if result.l_firstname is not None else results2.AttorneyAgentFirstName
                },
                {
                    "FieldName": "1lawyer.MName",
                    "FieldValue": result.l_middlename if result.l_middlename is not None else results2.AttorneyAgentMiddleInitial
                },
                {
                    "FieldName": "1lawfirm.O.Addr123",
                    "FieldCalcOverride": True,
                    "FieldValue": result.l_address1 if result.l_address1 is not None else results2.AttorneyAgentAddress1
                },
                {
                    "FieldName": "1lawfirm.O.AddrUnitType",
                    "FieldValue": AptSteFlr(result.l_addressunittype)
                },
                {
                    "FieldName": "1lawfirm.O.Addr4",
                    "FieldValue": result.l_addressunitnumber if result.l_addressunitnumber is not None else results2.AttorneyAgentAddress2
                },
                {
                    "FieldName": "1lawfirm.O.City",
                    "FieldValue": result.l_city if result.l_city is not None else results2.AttorneyAgentCity
                },
                {
                    "FieldName": "1lawfirm.O.State",
                    "FieldValue": result.l_state if result.l_state is not None else results2.AttorneyAgentState
                },
                {
                    "FieldName": "1lawfirm.O.Zip",
                    "FieldValue": result.l_zipcode if result.l_zipcode is not None else results2.AttorneyAgentZipCode
                },
                {
                    "FieldName": "1lawfirm.O.Country",
                    "FieldValue": result.l_country if result.l_country is not None else results2.AttorneyAgentCountry
                },
                {
                    "FieldName": "1lawyer.O.Phone",
                    "FieldValue": result.l_phonenumber if result.l_phonenumber is not None else results2.AttorneyAgentPhoneNumber
                },
                {
                    "FieldName": "1lawyer.O.Mobile",
                    "FieldValue": iff(result.l_mobilenumber)
                },
                {
                    "FieldName": "1lawyer.O.Email",
                    "FieldValue": result.l_email if result.l_email is not None else results2.AttorneyAgentEmail
                },
                {
                    "FieldName": "1lawyer.LicNum",
                    "FieldValue": iff(result.l_BarNumber)
                },
                {
                    "FieldName": "1lawfirm.Company",
                    "FieldValue": result.l_firmname if result.l_firmname is not None else results2.AttorneyAgentFirmName
                },
                {
                    "FieldName": "USIMM82463.txt1.21.0",
                    "FieldValue": ('Form I-140 - ' + str(
                        result.b_lastname if result.b_lastname is not None else results2.BeneficiaryLastName) + ', ' + str(
                        result.b_firstname if result.b_firstname is not None else results2.BeneficiaryFirstName) + ' ' + str(
                        result.b_middlename if result.b_middlename is not None else results2.BeneficiaryMiddleName)).strip()
                },
                {
                    "FieldName": "1ben.LName",
                    "FieldValue": iff(result.LastName)
                },
                {
                    "FieldName": "1ben.FName",
                    "FieldValue": iff(result.FirstName)
                },
                {
                    "FieldName": "1ben.MName",
                    "FieldValue": iff(result.MiddleName)
                },
                {
                    "FieldName": "1ben.SEmp.Company",
                    "FieldValue": result.PetitionerName if result.PetitionerName is not None else results2.PetitionerName
                },
                {
                    "FieldName": "1ben.SignTitle",
                    "FieldValue": result.JobTitle if result.JobTitle is not None else results2.JobTitle
                },
                {
                    "FieldName": "1ben.M.Phone",
                    "FieldValue": result.PhoneNumber if result.PhoneNumber is not None else results2.PetitionerPhone
                },
                {
                    "FieldName": "1ben.M.Mobile",
                    "FieldValue": iff(result.MobilePhone)
                },
                {
                    "FieldName": "1ben.M.Email",
                    "FieldValue": iff(result.Email)
                },
                {
                    "FieldName": "1ben.M.Addr123",
                    "FieldCalcOverride": True,
                    "FieldValue": result.Address1 if result.Address1 is not None else results2.PetitionerAddress1
                },
                {
                    "FieldName": "1ben.M.Addr4",
                    "FieldValue": iff(result.AddressUnitNumber)
                },
                {
                    "FieldName": "1ben.M.City",
                    "FieldValue": result.City if result.City is not None else results2.PetitionerCity
                },
                {
                    "FieldName": "1ben.M.State",
                    "FieldValue": result.StateProvince if result.StateProvince is not None else results2.PetitionerState
                },
                {
                    "FieldName": "1ben.M.Zip",
                    "FieldValue": result.PostalZipCode if result.PostalZipCode is not None else results2.PetitionerZipCode
                },
                {
                    "FieldName": "1ben.M.Country",
                    "FieldValue": result.Country if result.Country is not None else results2.PetitionerPhone
                },

                {
                    "FieldName": "User.D937.AttorneyEligiblePracticeLawIn",
                    "FieldValue": 1
                },
                {
                    "FieldName": "User.D937.AmNotAmSubjectToAnyOrder",
                    "FieldValue": 1
                },
                {
                    "FieldName": "USIMM82463.chk1.02",
                    "FieldValue": 1
                },
                {
                    "FieldName": "USIMM82463.chk1.03",
                    "FieldValue": 2
                },
                {
                    "FieldName": "1altclient.EO.AddrUnitType",
                    "FieldValue": AptSteFlr(result.AddressUnitType)
                },
                {
                    "FieldName": "1lawyer.O.Fax",
                    "FieldValue": iff(result.l_Faxnumber)
                },
                {
                    "FieldName": "USIMM82463.chk1.25",
                    "FieldValue": 1
                },
                {
                    "FieldName": "1lawyer.PracticeJurisidiction",
                    "FieldValue": iff(result.l_LicensingAuthority)
                }
            ]

            params = {
                "HostFormOnQuik": True,
                "FormFields": data_dict2,
                "QuikFormID": "82463",
                "PrintEditablePDF": True
            }

            data_json = json.dumps(params)
            response = requests.post('https://websvcs.quikforms.com/rest/quikformsengine/qfe/execute/pdf',
                                     headers=headersAPI, data=data_json)
            api_response = response.json()
            pdf_base64 = api_response['ResultData']['PDF']
            # print()
            with open('Form G-28_for_I-140.pdf', 'wb') as pdf:
                pdf.write(base64.b64decode(pdf_base64))
            data_dict = {
                'organization_xref': result.OrganizationXref,
                'organization_name': result.OrganizationName,
                'petitioner_xref': result.PetitionerXref,
                'petitioner_name': result.PetitionerName,
                'beneficiary_xref': result.BeneficiaryXref,
                'project_xref': result.ProjectXref,
                'last_name': result.b_lastname,
                'first_name': result.b_firstname,
            }
            data_dict = {i: v.strip() for i, v in data_dict.items()}

    pdf_merge_sequence(data_dict, formtype, result, results2, most_recent_entry_rec, most_recent_notice, parent_dir,
                       iff, change_date, qpdf_merger)


def pdf_merge_sequence(data_dict, formtype, result, results2, most_recent_entry_rec, most_recent_notice, parent_dir,
                       iff, change_date, qpdf_merger):
    os.chdir(parent_dir)
    if data_dict['organization_xref']:
        target_dir = op.join(parent_dir,
                             'ImmiLytics',
                             f"{data_dict['organization_xref']} - {data_dict['organization_name']}",
                             f"{data_dict['petitioner_xref']} - {data_dict['petitioner_name']}",
                             f"{data_dict['beneficiary_xref']} - {data_dict['last_name']}, {data_dict['first_name']}",
                             f"{data_dict['project_xref']}")

    else:
        target_dir = op.join(parent_dir,
                             'ImmiLytics',
                             f"{data_dict['petitioner_xref']} - {data_dict['petitioner_name']}",
                             f"{data_dict['beneficiary_xref']} - {data_dict['last_name']}, {data_dict['first_name']}",
                             f"{data_dict['project_xref']}")

    support_doc_dir = op.join(target_dir, 'Supporting Docs')
    project_level_dir = op.dirname(target_dir)

    if not op.exists(support_doc_dir):
        Path(support_doc_dir).mkdir(parents=True, exist_ok=True)
    if op.exists(support_doc_dir):
        pdfs = []
        sequence = ["Masters Degree", "Masters Transcripts", "Bachelors Degree", "Bachelors Transcripts",
                    "EVL", "Additional Supporting Docs", "Misc. Docs", "Passport", "I-94", "Pay Stub",
                    "PR Sponsorship Letter"]
        sequence_multiple = [
            "Additional Supporting Docs", "Misc. Docs", "Pay Stubs"]
        pdf_files = glob.glob(op.join(support_doc_dir, "*.pdf"))
        shutil.move(op.join(parent_dir, 'Form G-28_for_I-140.pdf'),
                    op.join(target_dir, 'Form G-28_for_I-140.pdf'))
        shutil.move(op.join(parent_dir, 'Form I-140.pdf'),
                    op.join(target_dir, 'Form I-140.pdf'))

        ben_salutation = ''
        if result.ben_Salutation:
            ben_salutation = result.ben_Salutation + '.'
        else:
            if result.ben_gender:
                ben_salutation = 'Mr.' if result.ben_gender.lower().strip(
                ) == 'male' else 'Ms.' if result.ben_gender.lower().strip() == 'female' else ''

        rep_dic = {
            "currentdatehere": date.today().strftime('%B %d, %Y'),
            'petitionernamehere': result.PetitionerName,
            "beneficiaryfullnamehere": f'{result.b_firstname} {result.b_middlename}, {result.b_lastname}',
            "beneficiarylastnamehere": result.b_lastname,
            "etajobtitlehere": results2.JobTitle,
            "authorizedsignatoryfullnamehere": f'{result.FirstName} ,{result.LastName}',
            "authorizedsignatoryjobtitlehere": result.JobTitle,
            "salutationhere": ben_salutation,
            "primaryprojectattorneyhere": f'{result.l_firstname} {result.l_middlename}, {result.l_lastname}',
            "jobtitlefromresourcetablehere": iff(result.l_JobTitle),
            'clientprojectdeschere': result.ServiceType,
            'eprpetnrnamehere': result.PetitionerName,
            'clientfullnamehere': f'{result.b_firstname} {result.b_middlename}, {result.b_lastname}',
            'businessdescriptionhere': result.p_BusinessInfo,
            'clientsalutationhere': ben_salutation,
            'clientlnamehere': result.b_lastname,
            'projectproposedjobtitlehere': results2.JobTitle,
            '11111': results2.JobTitle,
            'heshehere': 'He' if result.ben_gender.lower().strip() == 'male' else 'She' if result.ben_gender.lower().strip() == 'female' else 'He/She',
            '22222': 'He' if result.ben_gender.lower().strip() == 'male' else 'She' if result.ben_gender.lower().strip() == 'female' else 'He/She',
            'clientprojectprimatyhere': f'{result.l_firstname} {result.l_middlename}, {result.l_lastname}',
            '33333': f'{result.l_firstname} {result.l_middlename}, {result.l_lastname}',
            'projectsponsorfnamehere': result.FirstName,
            '44444': result.FirstName,
            'projectsponsorlnamehere': result.LastName,
            '55555': result.LastName,
            'projectsponsortitlehere': result.JobTitle,
            '66666': result.JobTitle,
            'projectpetitionnamehere': result.ProjectPetitionName,
            'projectprimaryattorneyhere': f'{result.l_firstname} {result.l_middlename}, {result.l_lastname}',
            'maxoutexpdatehere': result.FinalNivDate.strftime('%m/%d/%Y') if result.FinalNivDate else '',
            'currentimmigrationstatusexpdatehere': change_date(most_recent_entry_rec.I94ExpirationDate),
            'i797expdatehere': change_date(most_recent_notice.I797ExpirationDate),
            'visaexpdatehere': change_date(result.vi_ValidTo),
            'eta9089jobtitlehere': iff(results2.JobTitle),
            'eta9089primaryworksitehere': iff(results2.WorkLocationAddress1),
            'eta9089address2here': iff(results2.WorkLocationAddress2),
            'eta9089cityhere': iff(results2.WorkLocationCity),
            'eta9089statehere': iff(results2.WorkLocationState) + ' - ' + str(iff(results2.WorkLocationZipCode)),
            'eta9089offeredwagefromhere': '$' + results2.OfferedWageFrom,
            'eta9089offeredwagetohere': '$' + results2.OfferedWageTo,
            'authorizedsignatoryfullnamehere': f'{result.FirstName} ,{result.LastName}',
            'authorizedsignatoryjobtitlehere': result.JobTitle
        }

        # mapping notes:
        # 11111-  projectproposedjobtitlehere
        # 22222 - heshehere
        # 33333 - clientprojectprimatyhere
        # 44444 - projectsponsorfnamehere
        # 55555 - projectsponsorlnamehere
        # 55555 - projectsponsortitlehere

        if op.exists(op.join(project_level_dir, 'I-140 Cover Letter.docx')):
            cover_doc = docx.Document(op.join(project_level_dir, 'I-140 Cover Letter.docx'))
            for para in cover_doc.paragraphs:
                for run in para.runs:
                    for search_word in rep_dic.keys():
                        if search_word in run.text:
                            run.text = run.text.replace(
                                search_word, rep_dic[search_word])

            cover_doc.save(op.join(target_dir, "I-140 Cover Letter.docx"))
            convert(op.join(target_dir, "I-140 Cover Letter.docx"),
                    op.join(target_dir, "I-140 Cover Letter.pdf"))

            pdfs.append(op.join(target_dir, 'I-140 Cover Letter.pdf'))

        pdfs.append(op.join(target_dir, 'Form G-28_for_I-140.pdf'))
        pdfs.append(op.join(target_dir, 'Form I-140.pdf'))
        pdfs.append(op.join(support_doc_dir, 'ETA 9089.pdf'))

        if op.exists(op.join(project_level_dir, 'I-140 - Employer Support Letter.docx')):
            support_doc = docx.Document(op.join(project_level_dir, 'I-140 - Employer Support Letter.docx'))

            for para in support_doc.paragraphs:
                for run in para.runs:
                    for search_word in rep_dic.keys():
                        if search_word in run.text:
                            run.text = run.text.replace(
                                search_word, rep_dic[search_word])

                if 'Insert job description from PERM' in para.text:
                    para.text = ''
                    a = para.add_run('Job Duties:')
                    a.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.add_run('\n\n')
                    para.add_run(
                        results2.JobDuties).alignment = WD_ALIGN_PARAGRAPH.LEFT

                    b = para.add_run(
                        '\n\nSpecific skills or other requirements:')
                    b.bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.add_run(
                        f'\n\n{results2.SpecialSkillsOtherRequirements}').alignment = WD_ALIGN_PARAGRAPH.LEFT

                if 'RE: ' in para.text:
                    for runs in para.runs:
                        runs.bold = True

            support_doc.save(op.join(target_dir, 'I-140 - Employer Support Letter.docx'))
            convert(op.join(target_dir, 'I-140 - Employer Support Letter.docx'),
                    op.join(target_dir, 'I-140 - Employer Support Letter.pdf'))
            pdfs.append(op.join(target_dir, 'I-140 - Employer Support Letter.pdf'))

        if op.exists(op.join(project_level_dir, "I-140 - Initial QC Doc.docx")):
            initial_qc_doc = docx.Document(op.join(project_level_dir, "I-140 - Initial QC Doc.docx"))
            for table in initial_qc_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for search_word in rep_dic.keys():
                            if search_word in cell.text:
                                cell.text = cell.text.replace(
                                    search_word, rep_dic[search_word])

            initial_qc_doc.save(op.join(target_dir, "I-140 - Initial QC Doc.docx"))

        for seq in sequence:
            for pdf_f in glob.glob(seq + "*"):
                pdfs.append(op.join(support_doc_dir, pdf_f))

        latest_annual_report = 0
        latest_annual_report_file = ''
        for pdf_f in glob.glob(op.join(project_level_dir, "*.pdf")):
            if "Annual Report" in pdf_f:
                file_name_annual = op.splitext(op.basename(pdf_f))
                file_name_split = (str(file_name_annual[0]).strip()).split(' ')
                if int(file_name_split[-1]) > latest_annual_report:
                    latest_annual_report = int(file_name_split[-1])
                    latest_annual_report_file = pdf_f
            elif "Tax Docs" in pdf_f:
                pdfs.append(op.join(project_level_dir, pdf_f))

        if latest_annual_report_file:
            pdfs.append(op.join(project_level_dir, latest_annual_report_file))

        os.chdir(parent_dir)
        qpdf_merger(pdfs, target_dir, formtype)

    else:
        print(f'Appropriate folder structure\n{support_doc_dir}\nnot found')
